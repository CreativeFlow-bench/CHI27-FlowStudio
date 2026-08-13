from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/experiments/flowstudio-participant-questionnaire-zh.md"
OUTPUT = ROOT / "docs/experiments/FlowStudio_参与者实验问卷_v0.9.docx"

INK = "172033"
BLUE = "2457C5"
LIGHT_BLUE = "EAF0FF"
LIGHT_GRAY = "F3F5F8"
MID_GRAY = "677085"
LINE = "CBD2DC"


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Arial Unicode MS"
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), "Arial Unicode MS")
    rfonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    rfonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_rich_text(paragraph, text, size=10.5, color=INK):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        set_font(paragraph.add_run(content), size=size, bold=bold, color=color)


def add_response_line(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_together = True
        set_font(p.add_run("________________________________________________________________________________"), 9, color=LINE)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("FLOWSTUDIO  |  PARTICIPANT STUDY"), 8, bold=True, color=MID_GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("参与者编号：____________    第 "), 8, color=MID_GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()._r
    run.append(fld_begin)
    run.append(instr)
    run.append(fld_end)
    set_font(p.add_run(" 页"), 8, color=MID_GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, before, after, color in (
        ("Title", 24, 0, 6, INK),
        ("Heading 1", 16, 14, 6, BLUE),
        ("Heading 2", 12.5, 10, 4, INK),
        ("Heading 3", 11, 7, 3, BLUE),
    ):
        style = styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc, lines):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_rich_text(title, lines[0].removeprefix("# "), size=24)

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    set_table_widths(meta, [3.45, 3.45])
    values = [
        "参与者编号：____________    顺序组：G1 / G2 / G3 / G4",
        "日期：____________",
        "研究员编号：____________",
        "开始：__________    结束：__________",
    ]
    for idx, cell in enumerate([c for row in meta.rows for c in row.cells]):
        shade(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        add_rich_text(p, values[idx], size=9.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(7)
    note.paragraph_format.space_after = Pt(8)
    add_rich_text(note, "填写说明：本问卷只使用参与者编号。没有正确或错误答案，请按照刚才的真实体验作答。", size=9.5, color=MID_GRAY)


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = [1.15, 2.65, 3.1] if len(rows[0]) == 3 else [7.0 / len(rows[0])] * len(rows[0])
    set_table_widths(table, widths)
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            if ri == 0:
                shade(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value, size=8.8, color=INK)
            for run in p.runs:
                run.bold = ri == 0
        if ri == 0:
            set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    b_start = next(i for i, line in enumerate(lines) if line.startswith("## B."))
    c_start = next(i for i, line in enumerate(lines) if line.startswith("## C."))
    d_start = next(i for i, line in enumerate(lines) if line.startswith("## D."))
    pre = lines[:b_start]
    b_block = lines[b_start:c_start]
    c_block = lines[c_start:d_start]
    tail = lines[d_start:]

    expanded = list(pre)
    for trial in range(1, 5):
        block = list(b_block)
        block[0] = f"## B{trial}. 第 {trial} 次任务后问卷"
        block = [line for line in block if "请将本页复制四份" not in line]
        expanded.extend(block)
    for condition in range(1, 3):
        block = list(c_block)
        block[0] = f"## C{condition}. 第 {condition} 个系统后问卷"
        expanded.extend(block)
    expanded.extend(tail)
    lines = expanded
    doc = Document()
    configure_document(doc)
    add_title_block(doc, lines)

    index = 1
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        index += 1
        if not stripped or stripped.startswith("**参与者编号") or stripped.startswith("**研究员编号") or stripped.startswith("> 本问卷"):
            continue

        if stripped.startswith("| "):
            table_lines = [stripped]
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("## "):
            heading = stripped[3:]
            if heading.startswith("B") or heading.startswith("C") or heading.startswith("D. 最终") or heading.startswith("E. 半") or heading.startswith("F. 研究"):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_rich_text(p, heading, size=16, color=BLUE)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_rich_text(p, stripped[4:], size=12.5)
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            shade_like = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_BLUE)
            shade_like.append(shd)
            add_rich_text(p, stripped[2:], size=9.5)
            continue

        is_numbered = bool(re.match(r"^\d+\. ", stripped))
        if is_numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(4)
            add_rich_text(p, stripped, size=10)
            if stripped.endswith("：________________________________") or stripped.endswith("：____________________________________"):
                p.paragraph_format.space_after = Pt(7)
            continue

        if stripped.startswith("□"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_rich_text(p, stripped, size=9.5)
            continue

        if stripped.startswith("________________________________________________________________"):
            add_response_line(doc)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_rich_text(p, stripped, size=9.8)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
